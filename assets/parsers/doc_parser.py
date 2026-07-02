"""
.doc 格式文件解析器
使用 python-docx 读取 .docx 文件（注意：.doc 旧格式兼容性有限）
对于真正的 .doc (非 .docx)，python-docx 可能无法读取，
这里尝试读取，失败时给出提示。
"""
import os


def parse_doc(file_path: str) -> list:
    """
    尝试解析 .doc / .docx 文件
    返回: list of rows (每行为单元格值列表)
    
    注意: 旧版 .doc 格式可能无法被 python-docx 解析，
    如果失败，会给出手动转换的建议。
    """
    try:
        from docx import Document
        doc = Document(file_path)
    except ImportError:
        raise RuntimeError("需要安装 python-docx: pip install python-docx")
    except Exception as e:
        raise RuntimeError(
            f"无法解析 .doc 文件 [{file_path}]: {e}\n"
            f"提示: 该文件可能是旧版 .doc 格式，建议用WPS/Office打开后另存为 .xlsx。"
        )

    # 尝试提取表格
    tables = doc.tables
    if not tables:
        # 无表格，尝试读取段落
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        return [{"paragraphs": paragraphs}]

    all_rows = []
    for table in tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            all_rows.append(cells)

    return all_rows


def parse_doc_as_flat(file_path: str) -> list:
    """
    将doc内容平坦化为列表(每行dict，含sheet/row_idx/row_data)
    文档只有一个隐性sheet "7期"
    """
    rows = parse_doc(file_path)
    result = []
    for idx, row in enumerate(rows):
        if isinstance(row, dict) and "paragraphs" in row:
            result.append({
                "sheet": "7期",
                "row_idx": idx,
                "row_data": row["paragraphs"],
                "is_paragraph": True,
            })
        else:
            result.append({
                "sheet": "7期",
                "row_idx": idx,
                "row_data": row,
                "is_paragraph": False,
            })
    return result
